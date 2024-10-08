import pandas as pd
import streamlit as st
# import openai
# from openai import OpenAI
from docx import Document
from io import BytesIO
import fitz  
import google.generativeai as genai
import base64


# Set your OpenAI API key here

# openai.api_key = r"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"
# openai.api_key = r"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"
data=[]



# def get_completion(prompt, model="gpt-4o-2024-05-13"):
            
#             messages = [{"role": "user", "content": prompt}]
#             response = openai.ChatCompletion.create(
#             model=model,
#             messages=messages,
#             temperature=0.2,
            
#             )
            # return response["choices"][0]["message"]["content"]

gemini_api_key = "AIzaSyBNKJ5UoqldD8BwNVCwbDHs3GquaZ9OIjM"
genai.configure(api_key=gemini_api_key)

# Set up the model with configuration
safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
]

generation_config = {
    "temperature": 0.8,
    "top_p": 0.9,
    "top_k": 50,
    "max_output_tokens": 200,
}



model = genai.GenerativeModel(model_name="gemini-pro", 
                              generation_config=generation_config,
                              safety_settings=safety_settings
                              )          

st.title("Resume Matcher")

# job_description= st.text_area("Enter Job Description")
# File upload for job description
job_description_file = st.file_uploader("Upload Job Description", type=["txt", "pdf", "docx"])

# File upload for resume
resume_files = st.file_uploader("Upload Your Resume", type=["txt", "pdf", "docx"],accept_multiple_files=True)

if st.button("Generate Report") and (job_description_file or job_description) is not None and resume_files is not None:
   
    # Read content from uploaded job description file
   
   
    # Check if it's a DOCX file
    if job_description_file is not None:
            job_description_content = job_description_file.read()
            if job_description_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(BytesIO(job_description_content))
                job_description = "\n".join(para.text for para in doc.paragraphs)
            elif job_description_file.type == "application/pdf":
                pdf = fitz.open(stream=job_description_content, filetype="pdf")
                job_description = ""
                for page_num in range(pdf.page_count):
                    page = pdf[page_num]
                    job_description += page.get_text()
            else:
                # Assume it's a text file
                job_description = job_description_content.decode("utf-8")


    # for resume_file in resume_files:
    i=0
    while i<len(resume_files):  
        # Read content from uploaded resume file
        resume_file=resume_files[i]
        resume_content = resume_file.read()
       
        # Check if it's a DOCX file
        if resume_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(BytesIO(resume_content))
            resume_input = "\n".join(para.text for para in doc.paragraphs)
        elif resume_file.type == "application/pdf":
            pdf = fitz.open(stream=resume_content, filetype="pdf")
            resume_input = ""
            for page_num in range(pdf.page_count):
                page = pdf[page_num]
                resume_input += page.get_text()
        else:
            # Assume it's a text file
            resume_input = resume_content.decode("utf-8")

        
        input_text = f""" Your task is to find given details :{job_description} form the resume and dont give any more detail only give ask details .
               this is resume: {resume_input} find the details above from this

            """
        # st.write(input_text)
        # st.write("$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$4")

        generated_resume = model.generate_content([input_text])
       

        data.append({"Name":resume_file.name ,"Summary-Match Percentage":str(generated_resume.text)})
        i=i+1
    df = pd.DataFrame(data)
    # st.write("scs")
    st.write(df)
   
    csv = df.to_csv(index=False)
    b64 = base64.b64encode(csv.encode()).decode()  # Encode to base64
    href = f'<a href="data:file/csv;base64,{b64}" download="data.csv"> Click On For Download CSV File</a>'
    st.markdown(href, unsafe_allow_html=True)

       
